# tools.py
import os
import pandas as pd
from docx import Document as PyDocxDocument
import re
import logging
import ast

from langchain.vectorstores import FAISS
from langchain_community.embeddings import OllamaEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import RetrievalQA
from langchain_community.llms import Ollama as LangChainOllama
from langchain_core.documents import Document as LangChainDocument

logger = logging.getLogger(__name__)

DOC_PATH = "policy.docx"
EMBEDDING_MODEL_NAME = "nomic-embed-text"
QA_LLM_MODEL = "mistral"

FAISS_INDEX_PATH = "faiss_policy_index"

POLICY_QA_CHAIN = None  # Global variable for the QA chain

# --- Tool Base Class ---
class Tool:
    def __init__(self, name, description):
        self.name = name
        self.description = description

    def execute(self, params_str: str):
        raise NotImplementedError


# --- LangChain Powered Q&A System Initialization ---
def initialize_policy_qa_system():
    global POLICY_QA_CHAIN
    if POLICY_QA_CHAIN is not None and not (isinstance(POLICY_QA_CHAIN, str) and POLICY_QA_CHAIN.startswith("ERROR:")):
        logger.info("Policy QA system already initialized and seems valid.")
        return True # Indicate success or already done

    logger.info(f"Attempting to initialize Policy QA system...")
    if not os.path.exists(DOC_PATH):
        error_msg = f"Error: Document '{DOC_PATH}' not found. Cannot initialize Policy Q&A system."
        logger.error(error_msg)
        POLICY_QA_CHAIN = f"ERROR: {error_msg}" # Set global error state
        return False

    try:
        OLLAMA_BASE_URL_FROM_ENV = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
        logger.info(f"PYTHON SCRIPT in tools.py is using OLLAMA_BASE_URL: {OLLAMA_BASE_URL_FROM_ENV}")
        
        logger.info(f"Initializing OllamaEmbeddings with model: {EMBEDDING_MODEL_NAME} using base_url: {OLLAMA_BASE_URL_FROM_ENV}")
        embeddings = OllamaEmbeddings(model=EMBEDDING_MODEL_NAME, base_url=OLLAMA_BASE_URL_FROM_ENV)

        vector_store = None
        if os.path.exists(FAISS_INDEX_PATH):
            try:
                logger.info(f"Loading existing FAISS index from {FAISS_INDEX_PATH}...")
                vector_store = FAISS.load_local(FAISS_INDEX_PATH, embeddings, allow_dangerous_deserialization=True)
                logger.info("FAISS index loaded successfully.")
            except Exception as load_exc:
                logger.warning(f"Could not load existing FAISS index (path: {FAISS_INDEX_PATH}): {load_exc}. Will re-create.")
                # vector_store remains None
        else:
            logger.info(f"No existing FAISS index found at {FAISS_INDEX_PATH}. Creating a new one.")
            # vector_store remains None

        if vector_store is None: # Create new index
            logger.info(f"Creating new FAISS index from {DOC_PATH}...")
            doc = PyDocxDocument(DOC_PATH)
            full_text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            if not full_text.strip():
                error_msg = f"Error: Document '{DOC_PATH}' is empty. Cannot build new index."
                logger.error(error_msg)
                POLICY_QA_CHAIN = f"ERROR: {error_msg}"
                return False

            lc_docs = [LangChainDocument(page_content=full_text, metadata={"source": DOC_PATH})]
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
            doc_chunks = text_splitter.split_documents(lc_docs)

            if not doc_chunks:
                error_msg = "Error: Document splitting resulted in no chunks. Cannot build vector store."
                logger.error(error_msg)
                POLICY_QA_CHAIN = f"ERROR: {error_msg}"
                return False
            logger.info(f"Document split into {len(doc_chunks)} chunks for new index.")

            vector_store = FAISS.from_documents(doc_chunks, embeddings)
            logger.info("New FAISS vector store created.")
            try:
                vector_store.save_local(FAISS_INDEX_PATH)
                logger.info(f"New FAISS index saved to {FAISS_INDEX_PATH}.")
            except Exception as save_exc:
                logger.error(f"Could not save new FAISS index to {FAISS_INDEX_PATH}: {save_exc}")
                # Continue with in-memory version if save fails, but log it as an issue.

        logger.info(f"Initializing LangChainOllama (for QA) with model: {QA_LLM_MODEL} using base_url: {OLLAMA_BASE_URL_FROM_ENV}")
        llm = LangChainOllama(model=QA_LLM_MODEL, temperature=0.05, base_url=OLLAMA_BASE_URL_FROM_ENV)

        POLICY_QA_CHAIN = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=vector_store.as_retriever(search_kwargs={"k": 3}), # Assuming vector_store is defined
            return_source_documents=False
        )
        logger.info("Policy QA system initialized and chain created successfully.")
        return True

    except Exception as e:
        error_msg = f"Critical error during Policy QA system initialization: {type(e).__name__} - {str(e)}"
        logger.exception("Exception during Policy QA system initialization:")
        POLICY_QA_CHAIN = f"ERROR: {error_msg}" # Set global error state
        return False


# --- Tool Definitions ---
class DocumentQATool(Tool):
    def __init__(self):
        super().__init__(
            name="PolicyDocumentQA",
            description=f"Answers questions specifically about the '{DOC_PATH}' document (employee leave policy). Input should be the user's question about the policy."
        )
        # Initialization is now handled externally by app.py calling initialize_policy_qa_system()

    def execute(self, question: str) -> str:
        if POLICY_QA_CHAIN is None or (isinstance(POLICY_QA_CHAIN, str) and POLICY_QA_CHAIN.startswith("ERROR:")):
            error_detail = POLICY_QA_CHAIN if POLICY_QA_CHAIN else "QA system not initialized."
            # Attempt to re-initialize once if called and not ready
            logger.warning(f"PolicyDocumentQA tool executed but QA system not ready. Attempting re-init. Status: {error_detail}")
            if not initialize_policy_qa_system(): # Returns True on success
                 return f"Error: Policy Q&A system could not be initialized. Check logs. Detail: {POLICY_QA_CHAIN}"
            # Check again
            if POLICY_QA_CHAIN is None or (isinstance(POLICY_QA_CHAIN, str) and POLICY_QA_CHAIN.startswith("ERROR:")):
                return f"Error: Policy Q&A system remains uninitialized after attempt. Detail: {POLICY_QA_CHAIN}"

        if not isinstance(POLICY_QA_CHAIN, RetrievalQA): # Should be the chain object itself
            return f"Error: Policy Q&A system has an invalid chain object type: {type(POLICY_QA_CHAIN)}"

        if not question.strip():
            return "Error: No question provided to PolicyDocumentQA tool."
        try:
            logger.info(f"Executing PolicyDocumentQA with question: {question}")
            result_dict = POLICY_QA_CHAIN.invoke({"query": question}) # LangChain expects dict input for query
            answer = result_dict.get("result", "Could not find a specific answer in the document for this query.")
            return f"Answer from '{DOC_PATH}': {answer}"
        except Exception as e:
            logger.exception(f"Error during PolicyDocumentQA execution for question '{question}':")
            return f"Error during PolicyDocumentQA execution: {type(e).__name__} - {str(e)}"


class FileReaderTool(Tool):
    def __init__(self):
        super().__init__(
            name="FileReaderTool",
            description="Reads content of 'report.csv' for sales data analysis or very simple .txt files. Input: filename (e.g., 'report.csv'). For policy document queries, use PolicyDocumentQA."
        )
        # This cache is mostly for the LLM to see what data FileReaderTool extracted,
        # rather than for inter-tool data passing in this design.
        self.extracted_data_cache = {}

    def execute(self, filename: str) -> str:
        processed_filename = filename.strip().replace("'", "").replace("\"", "")
        if not processed_filename:
            return "Error: No filename provided to FileReaderTool."
        if processed_filename.lower() != "report.csv" and not processed_filename.lower().endswith(".txt"):
            return f"Error: FileReaderTool is primarily for 'report.csv' or .txt files. For '{DOC_PATH}', use PolicyDocumentQA."

        abs_filepath = os.path.abspath(processed_filename)
        if not os.path.exists(abs_filepath) or not os.path.isfile(abs_filepath):
            return f"Error: File '{processed_filename}' not found or is not a file at '{abs_filepath}'."

        try:
            if processed_filename.lower() == "report.csv":
                df = pd.read_csv(abs_filepath)
                if df.empty:
                    return f"File '{processed_filename}' (CSV) is empty."
                columns_info = ", ".join(df.columns)
                data_preview = df.head(3).to_string(index=False)
                # Example of extracting specific numerical data for the LLM to see and use with CalculatorTool
                if 'Revenue' in df.columns and pd.api.types.is_numeric_dtype(df['Revenue']):
                    revenue_values = df['Revenue'].dropna().tolist()
                    self.extracted_data_cache[processed_filename + "_revenue"] = revenue_values # Example caching
                    return (f"Successfully read '{processed_filename}' (CSV). Columns: {columns_info}. "
                            f"Extracted 'Revenue' values: {revenue_values}. "
                            f"Use CalculatorTool with these values if calculation is needed. Data preview:\n{data_preview}")
                else:
                    return (f"Successfully read '{processed_filename}' (CSV). Columns: {columns_info}. "
                            f"Data preview:\n{data_preview}. No specific numerical 'Revenue' column for auto-extraction, but other calculations might be possible on numbers observed in preview.")
            elif processed_filename.lower().endswith(".txt"):
                with open(abs_filepath, 'r', encoding='utf-8') as f:
                    content = f.read(1000) # Snippet for .txt
                return f"Successfully read snippet from '{processed_filename}' (TXT):\n{content[:500]}..."
            else: # Should be caught by earlier check
                return f"Error: Unsupported file type for FileReaderTool: '{processed_filename}'."
        except Exception as e:
            logger.exception(f"Error processing file '{processed_filename}' with FileReaderTool:")
            return f"Error processing file '{processed_filename}': {type(e).__name__} - {str(e)}."


class CalculatorTool(Tool):
    def __init__(self):
        super().__init__(
            name="CalculatorTool",
            description="Evaluates mathematical expressions (e.g., '2+2*4') or sums lists of numbers provided as a string (e.g., '[10000, 12000]')."
        )

    def execute(self, expression_str: str) -> str:
        expression_str = expression_str.strip()
        try:
            if expression_str.startswith('[') and expression_str.endswith(']'):
                numbers = ast.literal_eval(expression_str)
                if not isinstance(numbers, list) or not all(isinstance(x, (int, float)) for x in numbers):
                    return "Error: Invalid list format or non-numerical content in list."
                return f"Result: {sum(numbers) if numbers else 0}"
            elif re.fullmatch(r"^[0-9.+\-*/()\s\.]+$", expression_str) and any(char.isdigit() for char in expression_str):
                result = eval(expression_str, {"__builtins__": {}}, {})
                return f"Result: {result}"
            else:
                return "Error: Invalid characters/structure. Use numbers, operators (+,-,*,/), parentheses, or list string like '[1,2,3]'."
        except ZeroDivisionError: return "Error: Division by zero."
        except SyntaxError: return f"Error: Invalid syntax in '{expression_str}'."
        except Exception as e:
            logger.exception(f"Error in CalculatorTool for '{expression_str}':")
            return f"Error evaluating '{expression_str}': {type(e).__name__} - {str(e)}"


# This dictionary instantiates the tools.
TOOLS_AVAILABLE = {
    "PolicyDocumentQA": DocumentQATool(),
    "FileReaderTool": FileReaderTool(),
    "CalculatorTool": CalculatorTool()
}

def get_tool_descriptions_for_prompt():
    return "\n".join([
        f"- {tool.name}: {tool.description}"
        for tool_name, tool in TOOLS_AVAILABLE.items()
    ])

logger.info("Tools module loaded. `initialize_policy_qa_system()` should be called by app.py.")